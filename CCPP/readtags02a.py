"""
Pilot program to explore reading tags of the file.
https://python-docx.readthedocs.io/en/latest/index.html

activate py36a2021
"""

from docx import Document

if __name__ == '__main__':

    document = Document('L000.docx')
    #print(dir(document))
    #print(document.settings.element.tag)
    #print(dir(document.core_properties))
    #print("\n\n")	
    #print(document.core_properties.category)
    #print(document.core_properties.identifier)
    print("keywords:\n *", document.core_properties.keywords)
    #print("document.paragraphs[0] = \n", document.paragraphs[0].text)
    for i, p in enumerate(document.paragraphs):
        print("parargraph[%d]:\n"%i, document.paragraphs[i].text)
    print("subject", document.core_properties.subject)
    print("title", document.core_properties.title)
	
	# It is not what (reading metadata tags) I want for now